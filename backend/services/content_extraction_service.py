"""
Wisdom Agent - Content Extraction Service

Extracts and cleans content from various sources:
- URLs (articles, blog posts, social media)
- Text (direct paste)
- Files (PDF, DOCX, TXT)

Author: Wisdom Agent Team
Date: 2025-12-20
Phase: 2, Day 5
"""

import logging
import re
from datetime import datetime
from typing import Optional, Dict, Any, Tuple
from urllib.parse import urlparse

import httpx
from bs4 import BeautifulSoup

# Optional imports - graceful degradation if not installed
try:
    from readability import Document as ReadabilityDocument
    HAS_READABILITY = True
except ImportError:
    HAS_READABILITY = False
    
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from backend.database.connection import get_db_session
from backend.database.fact_check_models import (
    ContentReview, SourceMetadata, SourceType
)

logger = logging.getLogger(__name__)


class ContentExtractionError(Exception):
    """Raised when content extraction fails."""
    pass


class PaywallDetectedError(ContentExtractionError):
    """Raised when a paywall is detected."""
    pass


class ContentExtractionService:
    """
    Service for extracting content from various sources.
    
    Handles:
    - URL fetching with proper headers and error handling
    - HTML parsing and article extraction
    - PDF text extraction
    - DOCX text extraction
    - Plain text handling
    - Metadata extraction (author, publication date, etc.)
    """
    
    # Common paywall indicators
    PAYWALL_INDICATORS = [
        "subscribe to continue",
        "subscription required",
        "premium content",
        "members only",
        "sign in to read",
        "create an account to continue",
        "this article is for subscribers",
        "unlock this article",
    ]
    
    # User agent for web requests
    USER_AGENT = (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    )
    
    def __init__(self):
        """Initialize the content extraction service."""
        self._http_client: Optional[httpx.AsyncClient] = None
    
    async def get_http_client(self) -> httpx.AsyncClient:
        """Get or create the HTTP client."""
        if self._http_client is None or self._http_client.is_closed:
            self._http_client = httpx.AsyncClient(
                timeout=30.0,
                follow_redirects=True,
                headers={"User-Agent": self.USER_AGENT}
            )
        return self._http_client
    
    async def close(self):
        """Close the HTTP client."""
        if self._http_client and not self._http_client.is_closed:
            await self._http_client.aclose()
    
    # ========================================================================
    # MAIN EXTRACTION METHOD
    # ========================================================================
    
    async def extract_content(self, review_id: int) -> Dict[str, Any]:
        """
        Extract content for a review based on its source type.
        
        Returns:
            Dict containing:
            - content: The extracted text content
            - title: Extracted or inferred title
            - metadata: Source metadata (author, date, etc.)
        """
        with get_db_session() as db:
            review = db.get(ContentReview, review_id)
            if not review:
                raise ContentExtractionError(f"Review {review_id} not found")
            
            source_type = review.source_type
            
            try:
                if source_type == SourceType.URL:
                    result = await self._extract_from_url(review.source_url)
                elif source_type == SourceType.TEXT:
                    result = await self._extract_from_text(review.source_content)
                elif source_type == SourceType.FILE:
                    result = await self._extract_from_file(review.source_content)
                else:
                    raise ContentExtractionError(f"Unknown source type: {source_type}")
                
                # Update review with extracted content
                review.source_content = result["content"]
                if result.get("title") and not review.title:
                    review.title = result["title"]
                
                # Save source metadata
                if result.get("metadata"):
                    self._save_metadata(db, review, result["metadata"])
                
                db.commit()
                
                logger.info(f"Extracted {len(result['content'])} chars for review {review_id}")
                return result
                
            except Exception as e:
                logger.exception(f"Content extraction failed for review {review_id}")
                raise ContentExtractionError(str(e))
    
    # ========================================================================
    # URL EXTRACTION
    # ========================================================================
    
    async def _extract_from_url(self, url: str) -> Dict[str, Any]:
        """
        Fetch and extract content from a URL.
        
        Uses readability-lxml for article extraction if available,
        falls back to BeautifulSoup.
        """
        logger.info(f"Fetching URL: {url}")
        
        # Validate URL
        parsed = urlparse(url)
        if not parsed.scheme or not parsed.netloc:
            raise ContentExtractionError(f"Invalid URL: {url}")
        
        # Fetch the page
        client = await self.get_http_client()
        try:
            response = await client.get(url)
            response.raise_for_status()
        except httpx.HTTPStatusError as e:
            if e.response.status_code == 404:
                raise ContentExtractionError(f"Page not found: {url}")
            elif e.response.status_code == 403:
                raise ContentExtractionError(f"Access forbidden: {url}")
            else:
                raise ContentExtractionError(f"HTTP error {e.response.status_code}: {url}")
        except httpx.RequestError as e:
            raise ContentExtractionError(f"Failed to fetch URL: {e}")
        
        html_content = response.text
        
        # Check for paywall
        if self._detect_paywall(html_content):
            raise PaywallDetectedError(
                f"Paywall detected at {url}. Please paste the article text instead."
            )
        
        # Extract article content
        if HAS_READABILITY:
            result = self._extract_with_readability(html_content, url)
        else:
            result = self._extract_with_beautifulsoup(html_content, url)
        
        # Extract metadata
        metadata = self._extract_url_metadata(html_content, url)
        result["metadata"] = metadata
        
        return result
    
    def _detect_paywall(self, html: str) -> bool:
        """Check if the page has a paywall."""
        html_lower = html.lower()
        for indicator in self.PAYWALL_INDICATORS:
            if indicator in html_lower:
                return True
        return False
    
    def _extract_with_readability(self, html: str, url: str) -> Dict[str, Any]:
        """Extract article using readability-lxml."""
        doc = ReadabilityDocument(html)
        
        # Get clean HTML and convert to text
        content_html = doc.summary()
        soup = BeautifulSoup(content_html, "html.parser")
        content = soup.get_text(separator="\n", strip=True)
        
        # Clean up whitespace
        content = self._clean_text(content)
        
        return {
            "content": content,
            "title": doc.title(),
        }
    
    def _extract_with_beautifulsoup(self, html: str, url: str) -> Dict[str, Any]:
        """Extract article using BeautifulSoup (fallback)."""
        soup = BeautifulSoup(html, "html.parser")
        
        # Remove script and style elements
        for element in soup(["script", "style", "nav", "header", "footer", "aside"]):
            element.decompose()
        
        # Try to find article content
        article = (
            soup.find("article") or 
            soup.find("main") or 
            soup.find(class_=re.compile(r"article|content|post|entry", re.I)) or
            soup.find("body")
        )
        
        if article:
            content = article.get_text(separator="\n", strip=True)
        else:
            content = soup.get_text(separator="\n", strip=True)
        
        # Get title
        title = None
        title_tag = soup.find("title")
        if title_tag:
            title = title_tag.get_text(strip=True)
        
        # Clean up
        content = self._clean_text(content)
        
        return {
            "content": content,
            "title": title,
        }
    
    def _extract_url_metadata(self, html: str, url: str) -> Dict[str, Any]:
        """Extract metadata from HTML page."""
        soup = BeautifulSoup(html, "html.parser")
        metadata = {
            "domain": urlparse(url).netloc,
        }
        
        # Try to find author
        author_meta = (
            soup.find("meta", {"name": "author"}) or
            soup.find("meta", {"property": "article:author"}) or
            soup.find("meta", {"name": "dc.creator"})
        )
        if author_meta and author_meta.get("content"):
            metadata["author"] = author_meta["content"]
        
        # Try to find publication date
        date_meta = (
            soup.find("meta", {"property": "article:published_time"}) or
            soup.find("meta", {"name": "publication_date"}) or
            soup.find("meta", {"name": "date"}) or
            soup.find("time", {"datetime": True})
        )
        if date_meta:
            date_str = date_meta.get("content") or date_meta.get("datetime")
            if date_str:
                try:
                    # Try common date formats
                    for fmt in ["%Y-%m-%dT%H:%M:%S", "%Y-%m-%d", "%Y/%m/%d"]:
                        try:
                            metadata["publish_date"] = datetime.strptime(
                                date_str[:len(fmt.replace("%", ""))], fmt
                            )
                            break
                        except ValueError:
                            continue
                except Exception:
                    pass
        
        # Try to find publication name
        site_name = soup.find("meta", {"property": "og:site_name"})
        if site_name and site_name.get("content"):
            metadata["publication"] = site_name["content"]
        
        return metadata
    
    # ========================================================================
    # TEXT EXTRACTION
    # ========================================================================
    
    async def _extract_from_text(self, text: str) -> Dict[str, Any]:
        """Process directly pasted text content."""
        content = self._clean_text(text)
        
        # Try to extract a title from the first line if it looks like a heading
        lines = content.split("\n")
        title = None
        if lines and len(lines[0]) < 200 and not lines[0].endswith("."):
            title = lines[0]
        
        return {
            "content": content,
            "title": title,
            "metadata": {},
        }
    
    # ========================================================================
    # FILE EXTRACTION
    # ========================================================================
    
    async def _extract_from_file(self, file_reference: str) -> Dict[str, Any]:
        """
        Extract content from a file.
        
        File reference format: "file:{file_id}" or direct file path
        """
        # Parse file reference
        if file_reference.startswith("file:"):
            file_id = file_reference[5:]
            file_path = await self._get_file_path(file_id)
        else:
            file_path = file_reference
        
        # Determine file type and extract
        if file_path.lower().endswith(".pdf"):
            return await self._extract_from_pdf(file_path)
        elif file_path.lower().endswith(".docx"):
            return await self._extract_from_docx(file_path)
        elif file_path.lower().endswith(".txt"):
            return await self._extract_from_txt(file_path)
        else:
            raise ContentExtractionError(f"Unsupported file type: {file_path}")
    
    async def _get_file_path(self, file_id: str) -> str:
        """Get the file path for a file ID from the database."""
        # TODO: Implement file lookup from your file storage system
        # For now, assume file_id is the path
        return file_id
    
    async def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a PDF file."""
        if not HAS_PDFPLUMBER:
            raise ContentExtractionError(
                "PDF extraction requires pdfplumber. Install with: pip install pdfplumber"
            )
        
        try:
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            content = "\n\n".join(text_parts)
            content = self._clean_text(content)
            
            return {
                "content": content,
                "title": None,
                "metadata": {"source_type": "pdf"},
            }
        except Exception as e:
            raise ContentExtractionError(f"Failed to extract PDF: {e}")
    
    async def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a DOCX file."""
        if not HAS_DOCX:
            raise ContentExtractionError(
                "DOCX extraction requires python-docx. Install with: pip install python-docx"
            )
        
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n\n".join(paragraphs)
            content = self._clean_text(content)
            
            # Try to get title from document properties
            title = None
            if doc.core_properties.title:
                title = doc.core_properties.title
            
            metadata = {"source_type": "docx"}
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
            
            return {
                "content": content,
                "title": title,
                "metadata": metadata,
            }
        except Exception as e:
            raise ContentExtractionError(f"Failed to extract DOCX: {e}")
    
    async def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from a plain text file."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            
            content = self._clean_text(content)
            
            return {
                "content": content,
                "title": None,
                "metadata": {"source_type": "txt"},
            }
        except Exception as e:
            raise ContentExtractionError(f"Failed to read text file: {e}")
    
    # ========================================================================
    # HELPER METHODS
    # ========================================================================
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        # Replace multiple whitespace with single space (but preserve newlines)
        text = re.sub(r"[^\S\n]+", " ", text)
        
        # Replace multiple newlines with double newline
        text = re.sub(r"\n{3,}", "\n\n", text)
        
        # Remove leading/trailing whitespace from each line
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)
        
        # Remove leading/trailing whitespace from entire text
        text = text.strip()
        
        return text
    
    def _save_metadata(
        self, 
        db, 
        review: ContentReview, 
        metadata: Dict[str, Any]
    ):
        """Save extracted metadata to the database."""
        source_metadata = SourceMetadata(
            review_id=review.id,
            author=metadata.get("author"),
            publication=metadata.get("publication"),
            publish_date=metadata.get("publish_date"),
            domain=metadata.get("domain"),
        )
        db.add(source_metadata)


# ============================================================================
# SINGLETON INSTANCE
# ============================================================================

_content_extraction_service: Optional[ContentExtractionService] = None


def get_content_extraction_service() -> ContentExtractionService:
    """Get or create the content extraction service instance."""
    global _content_extraction_service
    if _content_extraction_service is None:
        _content_extraction_service = ContentExtractionService()
    return _content_extraction_service
